"""ממיר את docs/tech-spec.md למסמך Word ניתן לעריכה, עם RTL תקין ב-Word עצמו.

הכללים שמיושמים כאן (לפי מה ש-Word דורש, ולא מה שדי בו לצפייה ב-LibreOffice):

1. כל פסקה שיש בה עברית מקבלת <w:bidi/> ויישור לימין. שורה שכולה לטינית
   מקבלת בסיס LTR ויישור לשמאל, כדי שלא תיתלה על השוליים הימניים.
2. בפסקה מעורבת עברית+לטינית **אין** לסמן <w:rtl/> על אף ריצה. Word הופך
   בכוח מספרים ולטינית שנקלעים לריצה מסומנת, ואז 3.10 נהפך ל-10.3.
3. <w:rtl/> מסומן רק על ריצות עבריות בפסקה שאין בה אות לטינית — שם הוא
   מה שמעגן נקודתיים או סימן שאלה בקצה הנכון.
4. כל ריצה מקבלת w:cs (גופן מורכב) ו-w:szCs (גודל מורכב); הדגשה דורשת גם
   w:b וגם w:bCs. בלי אלה הגופן והגודל פשוט לא חלים על העברית.
5. טבלאות ותאים הם "סיפור" נפרד: הטבלה מקבלת <w:bidi/> ב-tblPr כדי
   שהעמודה הראשונה תשב מימין, וכל פסקה בתא עוברת את אותה לוגיקה.
6. אין get_display ואין תווי כיווניות (U+2066-2069) — Word מצייר אותם
   כריבועים.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

HEB = re.compile(r"[֐-׿יִ-ﭏ]")
LIST_MARKER = re.compile(r"^\d{1,2}\.$")

BODY_FONT = "David"
MONO_FONT = "Consolas"
ACCENT = RGBColor(0x8A, 0x6A, 0x2B)   # זהב עמום, כמו ערכת הצבעים של האתר
MUTED = RGBColor(0x55, 0x55, 0x55)


# ---------------------------------------------------------------- פיצול לפי כתב

def _strong(ch: str):
    """עברית → True, לטינית או ספרה → False, ניטרלי → None.

    ספרה נחשבת LTR בכוונה, כדי שמספר לא ייכנס לתוך ריצה עברית — Word הופך
    בכוח מספר שנתפס בריצה מסומנת ב-rtl.
    """
    if HEB.match(ch):
        return True
    if ch.isascii() and ch.isalnum():
        return False
    return None


def split_by_script(text: str):
    default_rtl = next((s for s in (_strong(c) for c in text) if s is not None), True)
    segments, buf, buf_rtl = [], "", None
    for ch in text:
        s = _strong(ch)
        kind = s if s is not None else (buf_rtl if buf_rtl is not None else default_rtl)
        if buf_rtl is None or kind == buf_rtl:
            buf, buf_rtl = buf + ch, kind
        else:
            segments.append((buf, buf_rtl))
            buf, buf_rtl = ch, kind
    if buf:
        segments.append((buf, buf_rtl))
    return segments


def shift_boundary_spaces(segments):
    """רווח בסוף ריצת LTR שלפני ריצת RTL עובר לתחילת ה-RTL.

    Word גוזם רווח בסוף ריצה בגבול כיוון, וכך «2.» נדבק לכותרת שאחריו.
    """
    out = [[s, r] for s, r in segments]
    for i in range(len(out) - 1):
        seg, rtl = out[i]
        nseg, nrtl = out[i + 1]
        if rtl is False and nrtl is True and seg.endswith(" "):
            stripped = seg.rstrip(" ")
            out[i][0] = stripped
            out[i + 1][0] = seg[len(stripped):] + nseg
    return [(s, r) for s, r in out if s]


def merge_list_marker(segments):
    if (len(segments) >= 2 and segments[0][1] is False
            and LIST_MARKER.match(segments[0][0].strip())
            and segments[1][1] is True):
        return [(segments[0][0] + segments[1][0], True)] + list(segments[2:])
    return list(segments)


def para_is_rtl(text: str) -> bool:
    if HEB.search(text):
        return True
    if any(c.isascii() and c.isalnum() for c in text):
        return False
    return True


# ------------------------------------------------------------- עיצוב תוך-שורתי

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def parse_inline(text: str):
    """→ [(טקסט, מודגש, מונוספייס)] . מפרק **הדגשה** ו-`קוד`."""
    parts = []
    for chunk in INLINE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**") and len(chunk) > 4:
            parts.append((chunk[2:-2], True, False))
        elif chunk.startswith("`") and chunk.endswith("`") and len(chunk) > 2:
            parts.append((chunk[1:-1], False, True))
        else:
            parts.append((chunk, False, False))
    return parts


def add_run(paragraph, text, *, font, size, bold, italic, rtl_ok, color=None):
    run = paragraph.add_run(text)
    rPr = run._r.get_or_add_rPr()
    # סדר הילדים ב-rPr חייב לעקוב אחרי הסכמה: rFonts, b, bCs, i, iCs, color, sz, szCs, rtl
    rPr.append(rPr.makeelement(qn("w:rFonts"), {
        qn("w:ascii"): font, qn("w:hAnsi"): font, qn("w:cs"): font}))
    if bold:
        rPr.append(rPr.makeelement(qn("w:b"), {}))
        rPr.append(rPr.makeelement(qn("w:bCs"), {}))
    if italic:
        rPr.append(rPr.makeelement(qn("w:i"), {}))
        rPr.append(rPr.makeelement(qn("w:iCs"), {}))
    if color is not None:
        run.font.color.rgb = color
    rPr.append(rPr.makeelement(qn("w:sz"), {qn("w:val"): str(int(size * 2))}))
    rPr.append(rPr.makeelement(qn("w:szCs"), {qn("w:val"): str(int(size * 2))}))
    if rtl_ok:
        rPr.append(rPr.makeelement(qn("w:rtl"), {}))
    return run


def fill_paragraph(paragraph, text, *, size=11, bold=False, italic=False,
                   font=BODY_FONT, color=None):
    """ממלא פסקה קיימת בטקסט מעוצב, עם כיווניות נכונה."""
    base_rtl = para_is_rtl(text)
    pPr = paragraph._p.get_or_add_pPr()
    if base_rtl:
        pPr.append(pPr.makeelement(qn("w:bidi"), {}))
    paragraph.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if base_rtl
                           else WD_ALIGN_PARAGRAPH.LEFT)
    # פסקה שיש בה ולו אות לטינית אחת נחשבת מעורבת: אף ריצה בה אינה מסומנת rtl.
    para_has_latin = any(c.isascii() and c.isalpha() for c in text)

    for chunk, chunk_bold, mono in parse_inline(text):
        seg_font = MONO_FONT if mono else font
        seg_size = size - 0.5 if mono else size
        segments = shift_boundary_spaces(merge_list_marker(split_by_script(chunk)))
        for seg, is_rtl in segments:
            add_run(paragraph, seg, font=seg_font, size=seg_size,
                    bold=bold or chunk_bold, italic=italic,
                    rtl_ok=bool(is_rtl and not para_has_latin), color=color)
    return paragraph


# ------------------------------------------------------------------- הרכבת מסמך

def set_section_rtl(doc):
    for section in doc.sections:
        sectPr = section._sectPr
        sectPr.append(sectPr.makeelement(qn("w:bidi"), {}))


def set_table_rtl(table):
    tblPr = table._tbl.tblPr
    tblPr.append(tblPr.makeelement(qn("w:bidi"), {}))


def add_paragraph(doc, text, *, style=None, **kw):
    p = doc.add_paragraph(style=style)
    return fill_paragraph(p, text, **kw)


def add_heading(doc, text, level):
    sizes = {1: 22, 2: 15, 3: 12.5}
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(20 if level <= 2 else 14)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    fill_paragraph(p, text, size=sizes.get(level, 12), bold=True,
                   color=ACCENT if level <= 2 else None)
    return p


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Pt(18)
        pf.right_indent = Pt(18)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        pPr = p._p.get_or_add_pPr()
        # בלוק קוד הוא תמיד LTR, גם כשיש בו הערה בעברית
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        shd = pPr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear", qn("w:fill"): "F4F1EA"})
        pPr.append(shd)
        for seg, is_rtl in split_by_script(line or " "):
            add_run(p, seg, font=MONO_FONT, size=9.5, bold=False,
                    italic=False, rtl_ok=False)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_md_table(doc, rows):
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=len(rows), cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    set_table_rtl(table)
    for r, row in enumerate(rows):
        for c in range(len(header)):
            cell = table.cell(r, c)
            text = row[c] if c < len(row) else ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            fill_paragraph(para, text or " ", size=9.5, bold=(r == 0))
            if r == 0:
                tcPr = cell._tc.get_or_add_tcPr()
                tcPr.append(tcPr.makeelement(qn("w:shd"), {
                    qn("w:val"): "clear", qn("w:fill"): "EFE9DC"}))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def convert(md_path: Path, out_path: Path):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.35
    set_section_rtl(doc)

    lines = md_path.read_text("utf-8").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            block, i = [], i + 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            add_code_block(doc, block)
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            rows, i = [], i
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = split_row(lines[i])
                if not all(re.fullmatch(r":?-{2,}:?", c or "") for c in cells if c):
                    rows.append(cells)
                i += 1
            if rows:
                width = max(len(r) for r in rows)
                add_md_table(doc, [r + [""] * (width - len(r)) for r in rows])
            continue

        if not stripped:
            i += 1
            continue

        if re.fullmatch(r"-{3,}", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            pPr = p._p.get_or_add_pPr()
            pbdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pbdr.makeelement(qn("w:bottom"), {
                qn("w:val"): "single", qn("w:sz"): "6",
                qn("w:space"): "1", qn("w:color"): "C9BFA6"})
            pbdr.append(bottom)
            pPr.append(pbdr)
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            add_heading(doc, stripped.lstrip("#").strip(), level)
            i += 1
            continue

        if stripped.startswith(">"):
            block = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                block.append(lines[i].strip().lstrip(">").strip())
                i += 1
            p = add_paragraph(doc, " ".join(block), size=10.5, italic=True,
                              color=MUTED)
            pf = p.paragraph_format
            pf.left_indent = Pt(20)
            pf.right_indent = Pt(20)
            continue

        bullet = re.match(r"^[*\-]\s+(.*)", stripped)
        number = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if bullet or number:
            body, buf = (bullet.group(1) if bullet else number.group(2)), []
            i += 1
            # שורות המשך של אותו פריט (הזחה, בלי סימן חדש)
            while (i < len(lines) and lines[i].startswith(("  ", "\t"))
                   and lines[i].strip()
                   and not re.match(r"^[*\-]\s|^\d+\.\s", lines[i].strip())):
                buf.append(lines[i].strip())
                i += 1
            text = " ".join([body] + buf)
            marker = "• " if bullet else f"{number.group(1)}. "
            p = add_paragraph(doc, marker + text, size=11)
            pf = p.paragraph_format
            pf.right_indent = Pt(0)
            pf.left_indent = Pt(0)
            pf.first_line_indent = Pt(0)
            pf.space_after = Pt(3)
            continue

        # פסקה רגילה — שורות רצופות עד שורה ריקה
        block = []
        while i < len(lines) and lines[i].strip() and not re.match(
                r"^\s*(#|\||>|```|[*\-]\s|\d+\.\s|-{3,}\s*$)", lines[i]):
            block.append(lines[i].strip())
            i += 1
        if block:
            add_paragraph(doc, " ".join(block))
        else:
            i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    src = Path(sys.argv[1])
    dst = Path(sys.argv[2])
    print("נכתב:", convert(src, dst))
